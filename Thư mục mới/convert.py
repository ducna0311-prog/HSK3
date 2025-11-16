import docx
import random
import json
import sys

def create_quiz_data(doc_path):
    print("Bat dau doc tep Word (theo yeu cau moi: lay Chu Han)...")

    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"--- LOI NGHIEM TRONG ---")
        print(f"Khong mo duoc tep: {doc_path}. Chi tiet: {e}")
        return

    if not doc.tables:
        print("--- LOI ---")
        print("Khong tim thay bang (table) nao trong tep Word.")
        return
        
    table = doc.tables[0]
    print(f"Da tim thay bang. Bat dau doc {len(table.rows)} hang...")
    
    all_data = []
    all_simplified_chars = [] # Danh sách chứa tất cả chữ Hán
    error_count = 0
    
    # Bỏ qua hàng tiêu đề (hàng 0)
    for i, row in enumerate(table.rows[1:], start=1):
        try:
            # --- THAY ĐỔI LOGIC ---
            # Cột 2 là Giản thể (index 1)
            # Cột 4 là Tiếng Việt (index 3)
            simplified = row.cells[1].text.strip() # Lấy Giản thể
            vietnamese = row.cells[3].text.strip() # Lấy Tiếng Việt
            
            if simplified and vietnamese:
                all_data.append({"question": vietnamese, "correct_answer": simplified})
                all_simplified_chars.append(simplified) # Thêm vào danh sách mồi
            else:
                error_count += 1
        
        except Exception as e:
            error_count += 1
            if error_count == 1:
                print(f"Phat hien hang loi dau tien o (khoang) hang {i+1}. Loi: {e}")
                print("Se bo qua hang nay va tiep tuc...")

    print("------------------------------------------")
    if error_count > 0:
        print(f"Da bo qua {error_count} hang (hang trong, loi, hoac o bi gop).")
        
    print(f"===> Da doc thanh cong {len(all_data)} tu vung. <===")
    
    quiz_data_list = []
    
    if not all_simplified_chars:
        print("--- LOI ---")
        print("Khong tim thay du lieu Chu Han (Gian the) nao.")
        print("Vui long kiem tra lai cot 2 trong tep Word.")
        return

    print(f"Bat dau tao {len(all_data)} cau hoi...")
    
    for item in all_data:
        question = item["question"]
        correct_answer = item["correct_answer"]
        
        # Tạo 3 đáp án sai ngẫu nhiên (lấy từ ds Chữ Hán)
        decoys = set()
        while len(decoys) < 3:
            random_char = random.choice(all_simplified_chars)
            if random_char != correct_answer:
                decoys.add(random_char)
        
        # Đảm bảo chúng ta có đủ 3 decoys
        while len(decoys) < 3:
            decoys.add(random.choice(all_simplified_chars))

        answers = [
            {"text": correct_answer, "correct": True},
            {"text": list(decoys)[0], "correct": False},
            {"text": list(decoys)[1], "correct": False},
            {"text": list(decoys)[2], "correct": False}
        ]
        random.shuffle(answers)
        
        quiz_data_list.append({
            "question": question,
            "answers": answers
        })

    output_filename = "quizData.js"
    try:
        with open(output_filename, "w", encoding="utf-8") as f:
            json_string = json.dumps(quiz_data_list, indent=2, ensure_ascii=False)
            f.write(f"const quizData = {json_string};")
        
        print("------------------------------------------")
        print(f"*** THANH CONG! (Voi dap an la Chu Han) ***")
        print(f"Da tao tep '{output_filename}' voi {len(quiz_data_list)} cau hoi.")
        
    except Exception as e:
        print(f"--- LOI KHI GHI TEP ---")
        print(f"Loi khi ghi tep '{output_filename}': {e}")

# --- Chạy chương trình ---
create_quiz_data("HSK3.docx")
